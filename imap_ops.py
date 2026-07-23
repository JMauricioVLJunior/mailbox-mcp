"""IMAP operations with a per-account connection pool and automatic retry.

Special folders (Sent/Trash/Drafts/Junk) are resolved per account:
explicit config wins, then RFC 6154 SPECIAL-USE flags, then common names.
"""

import datetime as dt
import imaplib
import re
import threading
import time

from imap_tools import AND, MailBox

import config
import semantics

_IDLE_RECHECK = 60  # seconds idle after which the connection is validated before reuse
_SOCKET_TIMEOUT = 60  # per socket operation; prevents a dead connection from hanging a tool call

_SPECIAL_CANDIDATES = {
    "sent": ["INBOX.Sent", "Sent", "Sent Items", "INBOX.Sent Items", "Sent Messages"],
    "trash": ["INBOX.Trash", "Trash", "Deleted Items", "INBOX.Deleted Items"],
    "drafts": ["INBOX.Drafts", "Drafts"],
    "junk": ["INBOX.Junk", "Junk", "Spam", "INBOX.Spam"],
}
_SPECIAL_CONFIG = {"sent": config.SENT_FOLDER, "trash": config.TRASH_FOLDER, "drafts": config.DRAFTS_FOLDER}


class _Entry:
    def __init__(self):
        self.mb = None
        self.lock = threading.Lock()
        self.last_used = 0.0
        self.special = None


_pool: dict = {}
_pool_lock = threading.Lock()


def _entry_for(imap: dict) -> _Entry:
    key = (imap["host"], imap["user"])
    with _pool_lock:
        entry = _pool.get(key)
        if entry is None:
            entry = _Entry()
            _pool[key] = entry
        return entry


def _fresh_login(imap: dict) -> MailBox:
    mb = MailBox(imap["host"], int(imap.get("port", 993)),
                 timeout=_SOCKET_TIMEOUT, ssl_context=config.ssl_context())
    mb.login(imap["user"], imap["password"])
    return mb


def _kill(entry: _Entry) -> None:
    if entry.mb is not None:
        try:
            entry.mb.client.shutdown()
        except Exception:
            pass
        entry.mb = None


def _ensure(entry: _Entry, imap: dict) -> MailBox:
    if entry.mb is not None and time.time() - entry.last_used > _IDLE_RECHECK:
        try:
            entry.mb.client.noop()
        except Exception:
            _kill(entry)
    if entry.mb is None:
        entry.mb = _fresh_login(imap)
    return entry.mb


def _resolve_special(mb) -> dict:
    detected, names = {}, []
    for f in mb.folder.list():
        names.append(f.name)
        flags = " ".join(f.flags).lower() if f.flags else ""
        for key in ("sent", "trash", "drafts", "junk"):
            if f"\\{key}" in flags and key not in detected:
                detected[key] = f.name

    resolved = {}
    for key, candidates in _SPECIAL_CANDIDATES.items():
        configured = _SPECIAL_CONFIG.get(key, "")
        if configured:
            resolved[key] = configured
        elif key in detected:
            resolved[key] = detected[key]
        else:
            # None (not a guessed name that may not exist) — callers that NEED the
            # folder raise a clear error via _require_special; queries keep working
            resolved[key] = next((c for c in candidates if c in names), None)
    return resolved


_SPECIAL_ENV = {"sent": "MCP_SENT_FOLDER", "trash": "MCP_TRASH_FOLDER", "drafts": "MCP_DRAFTS_FOLDER"}


def _require_special(sp: dict, key: str) -> str:
    folder = sp.get(key)
    if not folder:
        env = _SPECIAL_ENV.get(key)
        raise ValueError(f"could not determine the {key} folder on this server"
                         + (f"; set {env} explicitly" if env else ""))
    return folder


def _run(imap: dict, fn, retry: bool = True):
    """Run fn(mb, special) on a pooled IMAP connection, retrying once on connection loss.
    Non-idempotent operations (move/append) pass retry=False: if the connection drops
    after the server executed but before the response arrived, a retry would duplicate
    the action — better to fail and let the agent/user decide."""
    entry = _entry_for(imap)
    with entry.lock:
        for attempt in (1, 2):
            try:
                mb = _ensure(entry, imap)
                if entry.special is None:
                    entry.special = _resolve_special(mb)
                result = fn(mb, entry.special)
                entry.last_used = time.time()
                return result
            except (imaplib.IMAP4.abort, ConnectionError, EOFError, OSError):
                _kill(entry)
                if attempt == 2 or not retry:
                    raise


def special_folders(imap: dict) -> dict:
    return _run(imap, lambda mb, sp: dict(sp))


def verify_login(host: str, port: int, user: str, password: str) -> bool:
    try:
        mb = MailBox(host, port, timeout=_SOCKET_TIMEOUT,
                     ssl_context=config.ssl_context()).login(user, password)
        mb.logout()
        return True
    except Exception:
        return False


def _msg_summary(msg, folder: str = None, snippet_len: int = 0) -> dict:
    tags = [f for f in (msg.flags or ()) if not f.startswith("\\")]
    entry = {
        "uid": msg.uid,
        "date": msg.date.isoformat() if msg.date else None,
        "from": msg.from_,
        "to": list(msg.to),
        "subject": msg.subject,
        "unread": "\\Seen" not in (msg.flags or ()),
        "tags": tags,
    }
    if folder is not None:
        entry["folder"] = folder
    if snippet_len:
        body = msg.text or msg.html or ""
        entry["snippet"] = " ".join(body.split())[:snippet_len]
    return entry


def _build_criteria(unread=None, from_contains=None, to_contains=None, subject_contains=None,
                     tag=None, since=None, before=None) -> dict:
    criteria = {}
    if unread:
        criteria["seen"] = False
    if from_contains:
        criteria["from_"] = from_contains
    if to_contains:
        criteria["to"] = to_contains
    if subject_contains:
        criteria["subject"] = subject_contains
    if tag:
        criteria["keyword"] = tag
    if since:
        criteria["date_gte"] = dt.date.fromisoformat(since)
    if before:
        criteria["date_lt"] = dt.date.fromisoformat(before)
    return criteria


# ---------- queries ----------

_STATUS_RE = re.compile(r'^"?(.+?)"?\s+\(MESSAGES (\d+) UNSEEN (\d+)\)$')


def _status_all(mb) -> list:
    """All-folder counts in ONE round trip via LIST RETURN STATUS when the server
    supports it (Dovecot does); falls back to one STATUS per folder."""
    try:
        typ, _ = mb.client._simple_command('LIST', '""', '*', 'RETURN (STATUS (MESSAGES UNSEEN))')
        _, lines = mb.client.response('STATUS')
        if typ == "OK" and lines:
            results = []
            for raw in lines:
                if raw is None:
                    continue
                line = raw.decode("utf-8", errors="replace") if isinstance(raw, bytes) else str(raw)
                m = _STATUS_RE.match(line.strip())
                if m:
                    results.append({"folder": m.group(1), "total": int(m.group(2)), "unread": int(m.group(3))})
            if results:
                return results
    except Exception:
        pass
    results = []
    for f in mb.folder.list():
        try:
            status = mb.folder.status(f.name, ["MESSAGES", "UNSEEN"])
        except Exception:
            continue
        results.append({"folder": f.name, "total": status.get("MESSAGES", 0),
                        "unread": status.get("UNSEEN", 0)})
    return results


def unread_summary(imap: dict) -> list:
    def fn(mb, sp):
        return [{"folder": r["folder"], "unread": r["unread"]}
                for r in _status_all(mb) if r["unread"] > 0]
    return _run(imap, fn)


def mailbox_stats(imap: dict) -> list:
    return _run(imap, lambda mb, sp: _status_all(mb))


def list_unread(imap: dict, limit_per_folder: int = 10, snippet_len: int = 0) -> list:
    def fn(mb, sp):
        folders = [r["folder"] for r in _status_all(mb) if r["unread"] > 0]
        results = []
        for folder in folders:
            mb.folder.set(folder)
            for msg in mb.fetch(AND(seen=False), limit=limit_per_folder, reverse=True,
                                 mark_seen=False, bulk=True):
                results.append(_msg_summary(msg, folder, snippet_len))
        return results
    return _run(imap, fn)


def list_folders(imap: dict) -> list:
    return _run(imap, lambda mb, sp: [f.name for f in mb.folder.list()])


def list_emails(imap: dict, folder: str, unread=False, from_contains=None, to_contains=None,
                 subject_contains=None, tag=None, since=None, before=None,
                 limit: int = 50, snippet_len: int = 0) -> list:
    criteria = _build_criteria(unread, from_contains, to_contains, subject_contains, tag, since, before)

    def fn(mb, sp):
        mb.folder.set(folder)
        query = AND(**criteria) if criteria else "ALL"
        return [_msg_summary(m, snippet_len=snippet_len)
                for m in mb.fetch(query, limit=limit, reverse=True, mark_seen=False, bulk=True)]
    return _run(imap, fn)


def search_emails(imap: dict, query: str = None, field: str = "any", tag: str = None,
                   since: str = None, before: str = None, folders: list = None,
                   limit_per_folder: int = 50, preview: bool = False) -> list:
    criteria = {}
    if tag:
        criteria["keyword"] = tag
    elif query:
        field_key = {"subject": "subject", "from": "from_", "body": "text", "any": "text"}[field]
        criteria[field_key] = query
    if since:
        criteria["date_gte"] = dt.date.fromisoformat(since)
    if before:
        criteria["date_lt"] = dt.date.fromisoformat(before)
    if not criteria:
        raise ValueError("provide query, tag, or a date range")

    def fn(mb, sp):
        target_folders = folders or [f.name for f in mb.folder.list()]
        results = []
        for folder in target_folders:
            try:
                mb.folder.set(folder)
                matches = list(mb.fetch(AND(**criteria), limit=limit_per_folder, reverse=True,
                                         mark_seen=False, bulk=True))
            except Exception:
                continue
            for msg in matches:
                results.append(_msg_summary(msg, folder, 200 if preview else 0))
        return results
    return _run(imap, fn)


def read_email(imap: dict, folder: str, uid: str, max_chars: int = 5000) -> dict:
    def fn(mb, sp):
        mb.folder.set(folder)
        msg = next(iter(mb.fetch(AND(uid=uid), mark_seen=False)), None)
        if msg is None:
            raise ValueError(f"uid={uid} not found in '{folder}'")
        body = msg.text or msg.html or ""
        return {
            "from": msg.from_,
            "to": list(msg.to),
            "cc": list(msg.cc),
            "date": msg.date.isoformat() if msg.date else None,
            "subject": msg.subject,
            "message_id": (msg.obj.get("Message-ID") or "").strip(),
            "tags": [f for f in (msg.flags or ()) if not f.startswith("\\")],
            "attachments": [{"filename": a.filename, "content_type": a.content_type,
                              "size": len(a.payload)} for a in msg.attachments],
            "body": body[:max_chars],
            "body_truncated": len(body) > max_chars,
        }
    return _run(imap, fn)


def _extract_pdf(payload: bytes, max_chars: int) -> dict:
    from io import BytesIO

    from pypdf import PdfReader
    reader = PdfReader(BytesIO(payload))
    parts, total = [], 0
    for i, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        parts.append(f"--- page {i} ---\n{text}")
        total += len(text)
        if total > max_chars:
            break
    content = "\n".join(parts)
    return {"content": content[:max_chars], "truncated": len(content) > max_chars,
            "pages": len(reader.pages)}


def _extract_docx(payload: bytes, max_chars: int) -> dict:
    from io import BytesIO

    from docx import Document
    doc = Document(BytesIO(payload))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    content = "\n".join(parts)
    return {"content": content[:max_chars], "truncated": len(content) > max_chars}


def _extract_xlsx(payload: bytes, max_chars: int) -> dict:
    from io import BytesIO

    from openpyxl import load_workbook
    wb = load_workbook(BytesIO(payload), read_only=True, data_only=True)
    sheet_names = [ws.title for ws in wb.worksheets]
    parts, total = [], 0
    for ws in wb.worksheets:
        parts.append(f"=== sheet: {ws.title} ===")
        for row in ws.iter_rows(values_only=True):
            line = "\t".join("" if c is None else str(c) for c in row).rstrip()
            if line:
                parts.append(line)
                total += len(line)
            if total > max_chars:
                break
        if total > max_chars:
            break
    wb.close()
    content = "\n".join(parts)
    return {"content": content[:max_chars], "truncated": len(content) > max_chars,
            "sheets": sheet_names}


def _extract_text(payload: bytes, max_chars: int) -> dict:
    for enc in ("utf-8", "latin-1"):
        try:
            text = payload.decode(enc)
            return {"content": text[:max_chars], "truncated": len(text) > max_chars}
        except (UnicodeDecodeError, AttributeError):
            continue
    return {"error": "could not decode as text"}


# Attachments come from untrusted senders and are parsed on a shared server: cap input
# size and the decompressed size of zip-based formats (docx/xlsx) to prevent one user's
# agent from OOM-killing the process for everyone with a decompression bomb.
_MAX_ATTACHMENT_BYTES = 25 * 1024 * 1024       # refuse to parse payloads larger than this
_MAX_UNCOMPRESSED_BYTES = 200 * 1024 * 1024    # zip-bomb guard for docx/xlsx
_MAX_EXTRACT_CHARS = 200_000                    # server-side ceiling on caller max_chars


def _zip_bomb_guard(payload: bytes) -> None:
    """docx/xlsx are ZIP archives; refuse ones whose members decompress to a huge total."""
    import zipfile
    from io import BytesIO
    try:
        with zipfile.ZipFile(BytesIO(payload)) as zf:
            total = sum(i.file_size for i in zf.infolist())
    except zipfile.BadZipFile:
        return  # not a real zip — the format parser will reject it with a clear error
    if total > _MAX_UNCOMPRESSED_BYTES:
        raise ValueError(f"attachment expands to ~{total // (1024 * 1024)} MB uncompressed; "
                         "refused as a possible decompression bomb")


def get_attachment(imap: dict, folder: str, uid: str, filename: str, max_chars: int = 20000) -> dict:
    TEXT_TYPES = ("application/json", "application/xml", "application/csv", "message/rfc822")
    TEXT_EXTS = (".txt", ".csv", ".json", ".xml", ".md", ".log", ".html", ".htm", ".ics", ".eml", ".sql", ".py", ".js")
    DOCX_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    XLSX_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    max_chars = max(0, min(max_chars, _MAX_EXTRACT_CHARS))

    def fn(mb, sp):
        mb.folder.set(folder)
        msg = next(iter(mb.fetch(AND(uid=uid), mark_seen=False)), None)
        if msg is None:
            raise ValueError(f"uid={uid} not found in '{folder}'")
        for a in msg.attachments:
            if a.filename != filename:
                continue
            base = {"filename": filename, "content_type": a.content_type, "size": len(a.payload)}
            if len(a.payload) > _MAX_ATTACHMENT_BYTES:
                return {**base, "error": f"attachment too large to read "
                                          f"({len(a.payload) // (1024 * 1024)} MB; limit "
                                          f"{_MAX_ATTACHMENT_BYTES // (1024 * 1024)} MB)"}
            name = filename.lower()
            try:
                if name.endswith(".pdf") or a.content_type == "application/pdf":
                    return {**base, **_extract_pdf(a.payload, max_chars)}
                if name.endswith(".docx") or a.content_type == DOCX_TYPE:
                    _zip_bomb_guard(a.payload)
                    return {**base, **_extract_docx(a.payload, max_chars)}
                if name.endswith(".xlsx") or a.content_type == XLSX_TYPE:
                    _zip_bomb_guard(a.payload)
                    return {**base, **_extract_xlsx(a.payload, max_chars)}
                if (a.content_type.startswith("text/") or a.content_type in TEXT_TYPES
                        or name.endswith(TEXT_EXTS)):
                    return {**base, **_extract_text(a.payload, max_chars)}
            except Exception as exc:
                return {**base, "error": f"extraction failed: {exc}"}
            return {**base, "error": "unsupported format for direct reading (image/zip/legacy xls...) "
                                       "- supported: pdf, docx, xlsx, csv, txt, json, xml, html, ics"}
        available = [a.filename for a in msg.attachments]
        raise ValueError(f"attachment '{filename}' not found; available: {available}")
    return _run(imap, fn)


def _base_subject(subject: str) -> str:
    s = subject or ""
    prefixes = ("re:", "res:", "fwd:", "fw:", "enc:")
    changed = True
    while changed:
        changed = False
        stripped = s.strip()
        for p in prefixes:
            if stripped.lower().startswith(p):
                s = stripped[len(p):]
                changed = True
                break
    return s.strip()


def get_thread(imap: dict, folder: str, uid: str, max_messages: int = 30,
                all_folders: bool = False) -> list:
    def fn(mb, sp):
        mb.folder.set(folder)
        msg = next(iter(mb.fetch(AND(uid=uid), mark_seen=False)), None)
        if msg is None:
            raise ValueError(f"uid={uid} not found in '{folder}'")
        base = _base_subject(msg.subject)
        if not base:
            return [_msg_summary(msg, folder, 200)]

        if all_folders:
            target = [f.name for f in mb.folder.list()]
        else:
            target = [f for f in dict.fromkeys([folder, sp["sent"], "INBOX"]) if f]

        results = []
        for fname in target:
            try:
                mb.folder.set(fname)
                for m in mb.fetch(AND(subject=base), mark_seen=False, bulk=True):
                    if _base_subject(m.subject) == base:
                        results.append(_msg_summary(m, fname, 200))
            except Exception:
                continue
        results.sort(key=lambda e: e["date"] or "")
        return results[:max_messages]
    return _run(imap, fn)


# ---------- intelligence ----------

_PENDING_FETCH_CAP = 200  # newest messages per folder considered when hunting stalled threads


def pending_replies(imap: dict, days_back: int = 14, folders: list = None, limit: int = 30) -> list:
    """Threads whose LAST message is from the other party (i.e. awaiting our reply)."""
    me = imap["user"].lower()
    since = dt.date.today() - dt.timedelta(days=days_back)

    def fn(mb, sp):
        if folders:
            target = list(folders)
        else:
            target = [f.name for f in mb.folder.list()
                      if f.name not in (sp["trash"], sp["drafts"], sp["junk"])]
        if sp["sent"] and sp["sent"] not in target:
            target.append(sp["sent"])

        threads = {}
        for folder in target:
            try:
                mb.folder.set(folder)
                msgs = list(mb.fetch(AND(date_gte=since), mark_seen=False, bulk=True,
                                      headers_only=True, limit=_PENDING_FETCH_CAP, reverse=True))
            except Exception:
                continue
            for m in msgs:
                if m.date is None:
                    continue
                base = _base_subject(m.subject).lower()
                if not base:
                    continue
                info = threads.setdefault(base, {"last": None, "subject": _base_subject(m.subject)})
                entry = {"date": m.date, "from": m.from_ or "", "folder": folder, "uid": m.uid}
                try:
                    if info["last"] is None or m.date > info["last"]["date"]:
                        info["last"] = entry
                except TypeError:
                    continue

        now = dt.datetime.now(dt.timezone.utc)
        results = []
        for info in threads.values():
            last = info["last"]
            if last is None:
                continue
            frm = last["from"].lower()
            if frm == me or last["folder"] == sp["sent"]:
                continue
            try:
                age_days = (now - last["date"]).days
            except TypeError:
                age_days = None
            results.append({
                "subject": info["subject"],
                "from": last["from"],
                "folder": last["folder"],
                "folder_meaning": semantics.folder_meaning(last["folder"], sp),
                "uid": last["uid"],
                "date": last["date"].isoformat(),
                "waiting_days": age_days,
                "automated": semantics.is_automated_sender(last["from"], info["subject"]),
            })
        # humans first (automated=False), oldest first
        results.sort(key=lambda e: (e["automated"], e["date"]))
        return results[:limit]
    return _run(imap, fn)


def contact_history(imap: dict, email: str, days_back: int = 365, limit: int = 20) -> dict:
    since = dt.date.today() - dt.timedelta(days=days_back)

    def fn(mb, sp):
        received, sent = [], []
        for folder in [f.name for f in mb.folder.list()]:
            try:
                mb.folder.set(folder)
                criteria = AND(to=email, date_gte=since) if folder == sp["sent"] else AND(from_=email, date_gte=since)
                for m in mb.fetch(criteria, mark_seen=False, bulk=True, headers_only=True):
                    if m.date is None:
                        continue
                    entry = {"date": m.date.isoformat(), "folder": folder, "uid": m.uid,
                             "subject": m.subject, "direction": "sent" if folder == sp["sent"] else "received"}
                    (sent if folder == sp["sent"] else received).append(entry)
            except Exception:
                continue
        together = sorted(received + sent, key=lambda e: e["date"])
        return {
            "email": email,
            "received_count": len(received),
            "sent_count": len(sent),
            "first_contact": together[0]["date"] if together else None,
            "last_contact": together[-1]["date"] if together else None,
            "recent": list(reversed(together[-limit:])),
        }
    return _run(imap, fn)


# ---------- actions ----------

def _as_uid_list(uids) -> list:
    if isinstance(uids, str):
        return [u.strip() for u in uids.split(",") if u.strip()]
    return list(uids)


def tag_emails(imap: dict, folder: str, uids, keyword: str, remove: bool = False) -> int:
    uid_list = _as_uid_list(uids)

    def fn(mb, sp):
        mb.folder.set(folder)
        mb.flag(uid_list, [keyword], not remove)
        return len(uid_list)
    return _run(imap, fn)


def move_emails(imap: dict, folder: str, uids, destination: str) -> int:
    uid_list = _as_uid_list(uids)

    def fn(mb, sp):
        mb.folder.set(folder)
        mb.move(uid_list, destination)
        return len(uid_list)
    return _run(imap, fn, retry=False)


def delete_emails(imap: dict, folder: str, uids) -> int:
    """Move to Trash (reversible) — never permanently deletes."""
    uid_list = _as_uid_list(uids)

    def fn(mb, sp):
        mb.folder.set(folder)
        mb.move(uid_list, _require_special(sp, "trash"))
        return len(uid_list)
    return _run(imap, fn, retry=False)


def mark_read(imap: dict, folder: str, uids, read: bool = True) -> int:
    uid_list = _as_uid_list(uids)

    def fn(mb, sp):
        mb.folder.set(folder)
        mb.flag(uid_list, ["\\Seen"], read)
        return len(uid_list)
    return _run(imap, fn)


def create_folder(imap: dict, name: str) -> str:
    def fn(mb, sp):
        mb.folder.create(name)
        return name
    return _run(imap, fn)


def append_to_sent(imap: dict, message_bytes: bytes) -> None:
    _run(imap, lambda mb, sp: mb.append(message_bytes, _require_special(sp, "sent"),
                                         dt=None, flag_set=["\\Seen"]), retry=False)


def append_draft(imap: dict, message_bytes: bytes) -> str:
    def fn(mb, sp):
        drafts = _require_special(sp, "drafts")
        mb.append(message_bytes, drafts, dt=None, flag_set=["\\Draft", "\\Seen"])
        return drafts
    return _run(imap, fn, retry=False)


def fetch_for_reply(imap: dict, folder: str, uid: str) -> dict:
    def fn(mb, sp):
        mb.folder.set(folder)
        msg = next(iter(mb.fetch(AND(uid=uid), mark_seen=False)), None)
        if msg is None:
            raise ValueError(f"uid={uid} not found in '{folder}'")
        return {
            "from": msg.from_,
            "to": list(msg.to),
            "cc": list(msg.cc),
            "subject": msg.subject,
            "message_id": (msg.obj.get("Message-ID") or "").strip(),
            "references": (msg.obj.get("References") or "").strip(),
            "date": msg.date.isoformat() if msg.date else None,
            "body": (msg.text or "")[:2000],
        }
    return _run(imap, fn)
